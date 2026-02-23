"""Tests for src/ingestion/loader.py"""

import pytest
from pathlib import Path
from src.ingestion.loader import DocumentLoader


@pytest.fixture
def loader():
    return DocumentLoader()


class TestLoad:
    def test_load_txt(self, loader, tmp_text_file):
        doc = loader.load(tmp_text_file)
        assert "Hello world." in doc["text"]
        assert doc["metadata"]["file_type"] == ".txt"
        assert doc["metadata"]["source"] == "test.txt"
        assert doc["metadata"]["page_count"] == 1

    def test_load_markdown(self, loader, tmp_markdown_file):
        doc = loader.load(tmp_markdown_file)
        assert "# Heading" in doc["text"]
        assert doc["metadata"]["file_type"] == ".md"

    def test_file_not_found_raises(self, loader):
        with pytest.raises(FileNotFoundError, match="Document not found"):
            loader.load("/nonexistent/file.txt")

    def test_unsupported_extension_raises(self, loader, tmp_path):
        f = tmp_path / "data.csv"
        f.write_text("a,b,c")
        with pytest.raises(ValueError, match="Unsupported file type"):
            loader.load(f)

    def test_accepts_path_string(self, loader, tmp_text_file):
        doc = loader.load(str(tmp_text_file))
        assert doc["text"]

    def test_metadata_has_absolute_path(self, loader, tmp_text_file):
        doc = loader.load(tmp_text_file)
        assert Path(doc["metadata"]["file_path"]).is_absolute()

    def test_load_empty_text_file(self, loader, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        doc = loader.load(f)
        assert doc["text"] == ""


class TestLoadDirectory:
    def test_loads_supported_files_only(self, loader, tmp_doc_directory):
        docs = loader.load_directory(tmp_doc_directory)
        sources = {d["metadata"]["source"] for d in docs}
        assert "a.txt" in sources
        assert "b.md" in sources
        assert "skip.jpg" not in sources

    def test_not_a_directory_raises(self, loader, tmp_text_file):
        with pytest.raises(NotADirectoryError):
            loader.load_directory(tmp_text_file)

    def test_empty_directory(self, loader, tmp_path):
        empty = tmp_path / "empty_dir"
        empty.mkdir()
        docs = loader.load_directory(empty)
        assert docs == []

    def test_continues_on_individual_file_error(self, loader, tmp_path):
        """If one file fails to load, others should still be loaded."""
        good = tmp_path / "good.txt"
        good.write_text("valid content")
        bad = tmp_path / "bad.pdf"
        bad.write_bytes(b"not a real pdf")  # will fail to parse

        docs = loader.load_directory(tmp_path)
        # At least the good file should load
        assert any(d["metadata"]["source"] == "good.txt" for d in docs)


class TestLoadPdf:
    def test_load_pdf_with_text(self, loader, tmp_path):
        """Create a PDF with actual extractable text using reportlab or pypdf."""
        from unittest.mock import patch, MagicMock

        pdf_path = tmp_path / "test.pdf"
        # Create a minimal PDF with pypdf, then mock the text extraction
        from pypdf import PdfWriter
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(pdf_path, "wb") as f:
            writer.write(f)

        doc = loader.load(pdf_path)
        assert doc["metadata"]["file_type"] == ".pdf"
        assert doc["metadata"]["source"] == "test.pdf"
        assert doc["metadata"]["page_count"] == 1

    def test_load_pdf_extracts_page_text(self, loader, tmp_path):
        """Verify that pages with extractable text are captured."""
        from unittest.mock import patch, MagicMock
        from pypdf import PdfWriter

        pdf_path = tmp_path / "with_text.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with open(pdf_path, "wb") as f:
            writer.write(f)

        # Mock PdfReader to return pages with text
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page one content."
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page]

        with patch("src.ingestion.loader.PdfReader", return_value=mock_reader):
            doc = loader.load(pdf_path)
            assert "Page one content" in doc["text"]
            assert doc["metadata"]["page_count"] == 1


class TestLoadDocx:
    def test_load_docx(self, loader, tmp_path):
        from docx import Document as DocxDocument

        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        doc.save(str(docx_path))

        result = loader.load(docx_path)
        assert "First paragraph" in result["text"]
        assert "Second paragraph" in result["text"]
        assert result["metadata"]["file_type"] == ".docx"
        assert result["metadata"]["page_count"] == 1


class TestSupportedExtensions:
    def test_supported_set(self):
        assert DocumentLoader.SUPPORTED_EXTENSIONS == {".txt", ".md", ".pdf", ".docx"}
